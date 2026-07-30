"""Phase 2: Test Planning and Estimation.

Rather than asking the LLM to invent scope numbers (the exact failure mode
documented in Phase 1/7's READMEs), effort estimation here starts from
*real, programmatically-counted* project metrics — actual scenario count,
page object count, requirement count — and only asks the LLM to reason
about person-days given those concrete numbers, and to draft the prose
sections of the test plan document.

Outputs (in output/): effort_estimation.md, Test_Plan.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from docx import Document

from common.docx_reader import get_section, get_use_case_items
from common.feature_reader import parse_feature
from common.llm import OllamaClient

DOCX_PATH = REPO_ROOT / "docs" / "AI Usecase.docx"
FEATURE_PATH = REPO_ROOT / "features" / "multiCityFlightBooking.feature"
PAGES_DIR = REPO_ROOT / "src" / "pages"
OUTPUT_DIR = Path(__file__).resolve().parent / "output"

# The estimate's actual numbers come from fixed programmatic ratios below, not
# the LLM — an earlier version asked the model to compute person-days itself
# and it produced nonsense arithmetic (e.g. "27 automated scenarios / 6 ...
# = 4.5 person-days per scenario" used to justify 120 person-days just to
# *execute* an automated suite, which runs in minutes). The model is only
# asked to write the one-line rationale for a number it's given.
ESTIMATION_RATIONALE_PROMPT = """You are a QA test lead. For the activity "{activity}", the \
estimated effort is {days} person-days, computed from: {basis}. Write exactly one sentence of \
rationale for a test plan reader, referencing the actual number given. Don't propose a \
different number."""

# Rough, explainable ratios — not industry-calibrated, just transparent and
# proportional to real scope so the total scales sensibly with project size.
# See compute_effort_days() for how these are applied.
FRAMEWORK_SETUP_BASE_DAYS = 5
REPORTING_FIXED_DAYS = 3

TEST_PLAN_SECTION_PROMPT = """You are a QA test lead writing a test plan document. Given the \
PROJECT FACTS below (real facts about an existing, already-built project — you are documenting \
it, not proposing something hypothetical), write the "{section}" section of a test plan. Write \
in plain prose/short bullets as appropriate, grounded only in the given facts — do not invent \
team names, dates, or tools not mentioned. Keep it to a focused paragraph or short bullet list, \
no headings (the heading is added separately)."""


def count_page_objects() -> list[str]:
    return sorted(p.stem for p in PAGES_DIR.glob("*.js") if p.stem != "BasePage")


def gather_scope_metrics() -> dict:
    scenarios = parse_feature(FEATURE_PATH)
    positive = [s for s in scenarios if s.polarity == "positive"]
    negative = [s for s in scenarios if s.polarity == "negative"]
    use_case_items = get_use_case_items(DOCX_PATH)
    page_objects = count_page_objects()
    return {
        "total_scenarios": len(scenarios),
        "positive_scenarios": len(positive),
        "negative_scenarios": len(negative),
        "use_case_requirements": len(use_case_items),
        "page_objects": page_objects,
        "stlc_phases_covered": 7,
    }


def metrics_to_text(metrics: dict) -> str:
    return (
        f"- Automated Gherkin scenarios: {metrics['total_scenarios']} "
        f"({metrics['positive_scenarios']} positive, {metrics['negative_scenarios']} negative)\n"
        f"- Use-case requirements (from requirement doc, items a-j): {metrics['use_case_requirements']}\n"
        f"- Page Objects implemented: {len(metrics['page_objects'])} ({', '.join(metrics['page_objects'])})\n"
        f"- STLC phases covered by this project: {metrics['stlc_phases_covered']} of 7\n"
        f"- Application under test: multi-city flight booking with seats, baggage, meals, "
        f"insurance/GST, fare review, and payment (MakeMyTrip)"
    )


def compute_effort_days(metrics: dict) -> list[dict]:
    n_scenarios = metrics["total_scenarios"]
    n_requirements = metrics["use_case_requirements"]
    n_pages = len(metrics["page_objects"])

    rows = [
        {
            "activity": "Requirement analysis",
            "days": round(0.5 * n_requirements, 1),
            "basis": f"0.5 person-day x {n_requirements} use-case requirements",
        },
        {
            "activity": "Test design (manual cases + automation scripts)",
            "days": round(0.75 * n_scenarios, 1),
            "basis": f"0.75 person-day x {n_scenarios} scenarios (manual case authoring + automation script)",
        },
        {
            "activity": "Framework setup",
            "days": round(FRAMEWORK_SETUP_BASE_DAYS + 1.0 * n_pages, 1),
            "basis": f"{FRAMEWORK_SETUP_BASE_DAYS} person-day base + 1 person-day x {n_pages} page objects",
        },
        {
            "activity": "Test execution & stabilization",
            "days": round(0.25 * n_scenarios, 1),
            "basis": f"0.25 person-day x {n_scenarios} scenarios (stabilization pass, not raw execution time)",
        },
        {
            "activity": "Reporting & documentation",
            "days": REPORTING_FIXED_DAYS,
            "basis": f"flat {REPORTING_FIXED_DAYS} person-days",
        },
    ]
    return rows


def build_effort_estimation(client: OllamaClient, metrics: dict) -> str:
    metrics_text = metrics_to_text(metrics)
    rows = compute_effort_days(metrics)

    lines = []
    for row in rows:
        rationale = client.ask(
            ESTIMATION_RATIONALE_PROMPT.format(activity=row["activity"], days=row["days"], basis=row["basis"]),
            "",
        ).strip()
        lines.append(f"- **{row['activity']}: {row['days']} person-days** — {rationale}")

    total = sum(row["days"] for row in rows)
    lines.append(f"\n**Total: {total} person-days** (~{round(total / 5, 1)} working weeks for a single engineer)")

    return (
        f"# Effort Estimation\n\n## Scope metrics (actual, not estimated)\n\n{metrics_text}\n\n"
        f"## Estimate\n\nPerson-day figures are computed from the scope metrics above via fixed, "
        f"transparent ratios (see `generate_test_plan.py::compute_effort_days`) — the LLM contributes "
        f"the rationale text per activity, not the arithmetic (see Design notes in README.md for why).\n\n"
        + "\n".join(lines) + "\n"
    )


def build_project_facts(metrics: dict) -> str:
    objectives = get_section(DOCX_PATH, "Project Objectives")
    return (
        f"Project: AI-Powered Test Automation Assistant capstone — multi-city flight booking "
        f"with ancillary services, on MakeMyTrip.\n\n"
        f"Stated objectives:\n{objectives}\n\n"
        f"Scope metrics:\n{metrics_to_text(metrics)}\n\n"
        f"Tech stack: Playwright (JavaScript) + Cucumber BDD + Page Object Model for the "
        f"automation framework; Excel (exceljs/openpyxl) for test data; "
        f"multiple-cucumber-html-reporter for HTML reports; a separate Python CLI chatbot "
        f"(Phase 7) using a local Ollama LLM for requirement Q&A and Selenium code generation.\n"
        f"Version control: GitHub."
    )


def add_section(document: Document, heading: str, body: str):
    document.add_heading(heading, level=1)
    for line in body.strip().splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith(("-", "*")):
            document.add_paragraph(line.lstrip("-* ").strip(), style="List Bullet")
        else:
            document.add_paragraph(line)


def build_test_plan_docx(client: OllamaClient, metrics: dict, effort_summary: str, path: Path):
    facts = build_project_facts(metrics)
    sections = [
        "Scope",
        "Test Objectives",
        "Test Approach / Strategy",
        "Resources and Tools",
        "Entry and Exit Criteria",
        "Risks and Mitigations",
    ]

    document = Document()
    title = document.add_heading("Test Plan — Multi-City Flight Booking Automation", level=0)
    document.add_paragraph("Generated with Gen AI assistance (Phase 2) from real project scope metrics.")

    for section in sections:
        print(f"  Drafting section: {section}...")
        body = client.ask(TEST_PLAN_SECTION_PROMPT.format(section=section), f"PROJECT FACTS:\n{facts}")
        add_section(document, section, body)

    document.add_heading("Timeline / Effort Estimate", level=1)
    for line in effort_summary.strip().splitlines():
        line = line.strip()
        if line and not line.startswith("#"):
            document.add_paragraph(line)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def main():
    print("Gathering real scope metrics from the project...")
    metrics = gather_scope_metrics()
    print(metrics_to_text(metrics))

    client = OllamaClient()

    print("\nEstimating effort via local LLM (grounded in the metrics above)...")
    effort_md = build_effort_estimation(client, metrics)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    (OUTPUT_DIR / "effort_estimation.md").write_text(effort_md, encoding="utf-8")

    print("\nDrafting Test Plan document...")
    build_test_plan_docx(client, metrics, effort_md, OUTPUT_DIR / "Test_Plan.docx")

    print(f"\nDone. Output written to {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()
